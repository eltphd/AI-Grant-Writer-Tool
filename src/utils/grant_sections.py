"""Grant sections management and template utilities.

This module handles the six core grant application sections with structured
data storage, template generation, and RAG-enabled retrieval.
"""

import json
import os
from datetime import datetime
from dataclasses import dataclass, asdict
from typing import Dict, List, Optional, Any
import markdown
from docx import Document
from docx.shared import Inches

@dataclass
class GrantSection:
    """Represents a single grant section."""
    id: str
    title: str
    content: str
    target_length: str
    description: str
    word_count: int
    status: str
    last_updated: Optional[str] = None

@dataclass
class GrantDocument:
    """Represents a complete grant document."""
    project_id: str
    user_id: str
    sections: Dict[str, GrantSection]
    total_words: int
    completion_percentage: float
    last_updated: str
    chat_summary: str

class GrantSectionManager:
    """Manages grant sections and templates."""
    
    # Core sections schema
    CORE_SECTIONS = {
        "exec_summary": {
            "id": "exec_summary",
            "title": "Executive Summary / Cover Sheet",
            "target_length": "250–400 words",
            "description": "Concise snapshot of the project: who you are, what you want to do, funding request, and expected impact.",
            "placeholder": "Describe your organization, the project goals, funding amount requested, and expected impact in 250-400 words..."
        },
        "org_background": {
            "id": "org_background",
            "title": "Organization Background & Capacity",
            "target_length": "300–500 words",
            "description": "Mission, history, leadership, key accomplishments, and infrastructure that prove you can deliver.",
            "placeholder": "Describe your organization's mission, history, leadership team, key accomplishments, and capacity to deliver this project..."
        },
        "need_statement": {
            "id": "need_statement",
            "title": "Statement of Need / Problem Analysis",
            "target_length": "400–600 words",
            "description": "Data-rich justification for why the project matters. Include quantitative evidence and community voices.",
            "placeholder": "Describe the problem or need your project addresses. Include data, statistics, and community input to justify the urgency..."
        },
        "project_design": {
            "id": "project_design",
            "title": "Project/Program Design & Objectives",
            "target_length": "600–800 words",
            "description": "The 'how.' Break into SMART objectives, activities, timelines, staffing, and partnerships.",
            "placeholder": "Detail your project design, SMART objectives, activities, timeline, staffing plan, and key partnerships..."
        },
        "evaluation_plan": {
            "id": "evaluation_plan",
            "title": "Evaluation & Impact Measurement",
            "target_length": "300–500 words",
            "description": "Outcomes, KPIs, data-collection methods, and reporting cadence.",
            "placeholder": "Describe your evaluation plan, key performance indicators, data collection methods, and reporting schedule..."
        },
        "budget_sustainability": {
            "id": "budget_sustainability",
            "title": "Budget & Sustainability Plan",
            "target_length": "250–400 words",
            "description": "Itemized budget, justification notes, and post-grant funding strategy.",
            "placeholder": "Provide your detailed budget breakdown, cost justifications, and sustainability plan for post-grant funding..."
        }
    }

    def __init__(self):
        self.documents: Dict[str, GrantDocument] = {}

    def create_grant_document(self, project_id: str, user_id: str) -> GrantDocument:
        """Create a new grant document with all six sections."""
        sections = {}
        
        for section_id, section_info in self.CORE_SECTIONS.items():
            sections[section_id] = GrantSection(
                id=section_id,
                title=section_info["title"],
                content="",
                target_length=section_info["target_length"],
                description=section_info["description"],
                word_count=0,
                status="empty",
                last_updated=None
            )
        
        document = GrantDocument(
            project_id=project_id,
            user_id=user_id,
            sections=sections,
            total_words=0,
            completion_percentage=0.0,
            last_updated=datetime.utcnow().isoformat(),
            chat_summary=""
        )
        
        self.documents[f"{project_id}_{user_id}"] = document
        return document

    def get_grant_document(self, project_id: str, user_id: str) -> Optional[GrantDocument]:
        """Get an existing grant document."""
        key = f"{project_id}_{user_id}"
        return self.documents.get(key)

    def update_section_from_chat(self, project_id: str, user_id: str, section_id: str, content: str) -> Dict[str, Any]:
        """Update a section with content from chat conversation."""
        document = self.get_grant_document(project_id, user_id)
        if not document:
            document = self.create_grant_document(project_id, user_id)
        
        if section_id in document.sections:
            section = document.sections[section_id]
            section.content = content
            section.word_count = len(content.split())
            section.last_updated = datetime.utcnow().isoformat()
            section.status = self._get_section_status(content)
            
            # Update document stats
            self._update_document_stats(document)
            
            return {
                "success": True,
                "section_id": section_id,
                "word_count": section.word_count,
                "status": section.status
            }
        
        return {"success": False, "error": "Section not found"}

    def update_chat_summary(self, project_id: str, user_id: str, summary: str) -> bool:
        """Update the chat summary for the document."""
        document = self.get_grant_document(project_id, user_id)
        if document:
            document.chat_summary = summary
            document.last_updated = datetime.utcnow().isoformat()
            return True
        return False

    def _get_section_status(self, content: str) -> str:
        """Determine the status of a section based on content."""
        if not content or content.strip() == "":
            return "empty"
        
        word_count = len(content.split())
        if word_count < 50:
            return "draft"
        elif word_count < 200:
            return "developing"
        else:
            return "complete"

    def _update_document_stats(self, document: GrantDocument):
        """Update document statistics."""
        total_words = sum(section.word_count for section in document.sections.values())
        complete_sections = sum(1 for section in document.sections.values() if section.status == "complete")
        
        document.total_words = total_words
        document.completion_percentage = (complete_sections / len(document.sections)) * 100
        document.last_updated = datetime.utcnow().isoformat()

    def get_document_stats(self, project_id: str, user_id: str) -> Dict[str, Any]:
        """Get document statistics."""
        document = self.get_grant_document(project_id, user_id)
        if not document:
            return {
                "total_words": 0,
                "complete_sections": 0,
                "completion_percentage": 0.0,
                "last_updated": None
            }
        
        complete_sections = sum(1 for section in document.sections.values() if section.status == "complete")
        
        return {
            "total_words": document.total_words,
            "complete_sections": complete_sections,
            "completion_percentage": document.completion_percentage,
            "last_updated": document.last_updated
        }

    def export_to_markdown(self, project_id: str, user_id: str) -> str:
        """Export grant document as markdown."""
        document = self.get_grant_document(project_id, user_id)
        if not document:
            return "# Grant Application\n\nNo content available."
        
        markdown_content = "# Grant Application\n\n"
        
        if document.chat_summary:
            markdown_content += f"## Conversation Summary\n\n{document.chat_summary}\n\n"
        
        for section_id, section in document.sections.items():
            if section.content:
                markdown_content += f"## {section.title}\n\n{section.content}\n\n"
        
        return markdown_content

    def export_to_docx(self, project_id: str, user_id: str) -> bytes:
        """Export grant document as DOCX."""
        document = self.get_grant_document(project_id, user_id)
        if not document:
            # Return empty document
            doc = Document()
            doc.add_heading('Grant Application', 0)
            doc.add_paragraph('No content available.')
            return doc.save()
        
        doc = Document()
        doc.add_heading('Grant Application', 0)
        
        if document.chat_summary:
            doc.add_heading('Conversation Summary', level=1)
            doc.add_paragraph(document.chat_summary)
            doc.add_paragraph()
        
        for section_id, section in document.sections.items():
            if section.content:
                doc.add_heading(section.title, level=1)
                doc.add_paragraph(section.content)
                doc.add_paragraph()
        
        return doc.save()

    def get_section_templates(self) -> Dict[str, Any]:
        """Get grant section templates."""
        return self.CORE_SECTIONS

    def auto_populate_from_chat(self, project_id: str, user_id: str, chat_messages: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Auto-populate sections based on chat conversation."""
        document = self.get_grant_document(project_id, user_id)
        if not document:
            document = self.create_grant_document(project_id, user_id)
        
        # Analyze chat messages to extract relevant information
        chat_content = " ".join([msg.get("message", "") for msg in chat_messages])
        
        # Simple keyword-based section mapping
        section_mappings = {
            "exec_summary": ["organization", "project", "funding", "impact", "summary"],
            "org_background": ["mission", "history", "leadership", "accomplishments", "capacity"],
            "need_statement": ["problem", "need", "data", "statistics", "community", "urgency"],
            "project_design": ["objectives", "activities", "timeline", "staffing", "partnerships"],
            "evaluation_plan": ["evaluation", "outcomes", "kpis", "data collection", "reporting"],
            "budget_sustainability": ["budget", "cost", "funding", "sustainability", "financial"]
        }
        
        populated_sections = 0
        
        for section_id, keywords in section_mappings.items():
            if section_id in document.sections:
                # Check if chat content contains relevant keywords
                relevant_content = self._extract_relevant_content(chat_content, keywords)
                if relevant_content:
                    self.update_section_from_chat(project_id, user_id, section_id, relevant_content)
                    populated_sections += 1
        
        # Update chat summary
        summary = self._generate_chat_summary(chat_messages)
        self.update_chat_summary(project_id, user_id, summary)
        
        return {
            "success": True,
            "populated_sections": populated_sections,
            "total_sections": len(document.sections)
        }

    def _extract_relevant_content(self, chat_content: str, keywords: List[str]) -> str:
        """Extract content relevant to specific keywords."""
        # Simple implementation - in production, use more sophisticated NLP
        relevant_sentences = []
        sentences = chat_content.split('.')
        
        for sentence in sentences:
            sentence_lower = sentence.lower()
            if any(keyword in sentence_lower for keyword in keywords):
                relevant_sentences.append(sentence.strip())
        
        return ". ".join(relevant_sentences) if relevant_sentences else ""

    def _generate_chat_summary(self, chat_messages: List[Dict[str, Any]]) -> str:
        """Generate a summary of the chat conversation."""
        if not chat_messages:
            return ""
        
        # Extract key topics from chat
        topics = []
        for msg in chat_messages:
            message = msg.get("message", "").lower()
            if any(topic in message for topic in ["organization", "project", "funding", "budget", "evaluation"]):
                topics.append(msg.get("message", "")[:100] + "...")
        
        if topics:
            return f"Key discussion topics: {' '.join(topics[:3])}"
        else:
            return "Conversation in progress..."

# Global instance
grant_section_manager = GrantSectionManager() 